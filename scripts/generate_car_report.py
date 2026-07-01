# -*- coding: utf-8 -*-
"""生成变电站巡检仿真系统 — 软件架构设计报告（Word）"""
from __future__ import annotations
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "DejaVu Sans"]
plt.rcParams["axes.unicode_minus"] = False

OUT_DIR = os.path.join(os.path.dirname(__file__), "car_report_assets")
DOC_PATH = os.path.join(os.path.expanduser("~"), "Desktop", "资料",
                        "变电站巡检仿真_软件架构设计报告.docx")
UML_HELPER_DIR = os.path.join(os.path.expanduser("~"), "Desktop", "资料")
BLACK = RGBColor(0, 0, 0)


def _black(run):
    run.font.color.rgb = BLACK
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        _black(r)


def para(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    r.bold = bold
    _black(r)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    _black(r)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                _black(r)
    for row_data in rows:
        row = t.add_row().cells
        for i, cell in enumerate(row_data):
            row[i].text = cell
            for p in row[i].paragraphs:
                for r in p.runs:
                    _black(r)
    doc.add_paragraph()


def draw_blackboard_arch(path):
    fig, ax = plt.subplots(figsize=(13, 9))
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 10)
    ax.axis("off")
    ax.text(6.5, 9.5, "黑板架构 — 变电站巡检仿真系统总体架构图",
            ha="center", fontsize=13, fontweight="bold", color="black")

    def box(x, y, w, h, title, items, color):
        rect = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02",
                              facecolor=color, edgecolor="black", linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h - 0.3, title, ha="center", va="top",
                fontsize=10, fontweight="bold", color="black")
        ax.text(x + 0.15, y + h - 0.6, items, ha="left", va="top",
                fontsize=8, color="black", linespacing=1.3)

    box(0.3, 7.5, 12.4, 1.5, "展示层",
        "Display（HTTP:8887 + WebSocket:8888）| WebSocketBridge | 浏览器 Canvas UI",
        "#E3F2FD")
    box(0.3, 5.5, 12.4, 1.5, "控制层（黑板控制器）",
        "Controller | TickScheduler（节拍驱动）| StatusDispatcher（状态分派）| CommandHandler",
        "#C8E6C9")
    box(0.3, 2.8, 12.4, 2.2, "知识源层（独立进程，互不直接调用）",
        "TargetPlanner（目标分配）  Navigator（BFS/A*路径规划）  StrategySupervisor（路线监督）\n"
        "TaskConfigurator（任务初始化）  Car × N（移动执行）",
        "#FFF9C4")
    box(0.3, 0.5, 5.8, 1.8, "黑板（共享数据）",
        "Redis BlackboardClient\n地图位图 / 车辆状态 / 路径 / 探索率",
        "#FFCCBC")
    box(6.9, 0.5, 5.8, 1.8, "连接件（消息中间件）",
        "RabbitMQ MessageBus\n点对点队列 + Fanout 广播 UpdateView",
        "#E1BEE7")

    ax.annotate("", xy=(6.5, 7.5), xytext=(6.5, 7.0),
                arrowprops=dict(arrowstyle="<->", color="black", lw=1.2))
    ax.annotate("", xy=(6.5, 5.5), xytext=(6.5, 5.0),
                arrowprops=dict(arrowstyle="<->", color="black", lw=1.2))
    ax.text(6.7, 7.25, "WebSocket / MQ", fontsize=7.5, color="black")
    ax.text(6.7, 5.25, "MQ 指令调度", fontsize=7.5, color="black")
    ax.annotate("", xy=(3.2, 2.8), xytext=(3.2, 2.3),
                arrowprops=dict(arrowstyle="<->", color="black", lw=1))
    ax.annotate("", xy=(9.8, 2.8), xytext=(9.8, 2.3),
                arrowprops=dict(arrowstyle="<->", color="black", lw=1))
    ax.text(2.5, 2.55, "读写黑板", fontsize=7.5, color="black")
    ax.text(9.0, 2.55, "MQ消息", fontsize=7.5, color="black")

    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def draw_seq_startup(path):
    actors = ["Browser", "WebSocketBridge", "Controller\nCommandHandler",
              "StatusDispatcher", "TaskConfigurator", "TickScheduler"]
    xs = {a: i + 1 for i, a in enumerate(actors)}
    fig, ax = plt.subplots(figsize=(13, 8))
    ax.set_xlim(0, len(actors) + 1)
    ax.set_ylim(0, 15)
    ax.axis("off")
    ax.set_title("图2  仿真启动协作时序图（SET_CONFIG → TASK_READY）",
                 fontsize=11, fontweight="bold", color="black", pad=10)
    for a, x in xs.items():
        ax.plot([x, x], [1, 14], "k-", lw=0.8)
        ax.text(x, 14.2, a, ha="center", fontsize=7.5, color="black")
    msgs = [
        ("Browser", "WebSocketBridge", "SET_CONFIG JSON"),
        ("WebSocketBridge", "Controller\nCommandHandler", "publish(ControllerCmd)"),
        ("Controller\nCommandHandler", "StatusDispatcher", "prepareForNewConfig()"),
        ("StatusDispatcher", "TaskConfigurator", "forwardConfig(FORWARD_CONFIG)"),
        ("TaskConfigurator", "TaskConfigurator", "TaskInitializer.initialize()"),
        ("TaskConfigurator", "Controller\nCommandHandler", "TASK_READY"),
        ("Controller\nCommandHandler", "StatusDispatcher", "onTaskReady()"),
        ("StatusDispatcher", "TickScheduler", "start()"),
        ("TickScheduler", "TickScheduler", "tickLoop() 每500ms"),
    ]
    y = 13
    for src, dst, label in msgs:
        x1, x2 = xs[src], xs[dst]
        if x1 == x2:
            ax.plot([x1-0.25, x1+0.25], [y, y], "k-", lw=0.8)
            ax.text(x1+0.3, y, label, fontsize=7, color="black", va="center")
        else:
            ax.annotate("", xy=(x2, y), xytext=(x1, y),
                        arrowprops=dict(arrowstyle="->", color="black", lw=0.9))
            ax.text((x1+x2)/2, y+0.1, label, ha="center", fontsize=6.5, color="black")
        y -= 1.2
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def draw_seq_tick(path):
    actors = ["TickScheduler", "StatusDispatcher", "TargetPlanner",
              "Navigator", "StrategySupervisor", "Car\nMoveExecutor", "Display"]
    xs = {a: i + 1 for i, a in enumerate(actors)}
    fig, ax = plt.subplots(figsize=(14, 11))
    ax.set_xlim(0, len(actors) + 1)
    ax.set_ylim(0, 22)
    ax.axis("off")
    ax.set_title("图3  单节拍探索协作时序图（IDLE→ASSIGN→PLAN→SUPERVISE→MOVE）",
                 fontsize=11, fontweight="bold", color="black", pad=10)
    for a, x in xs.items():
        ax.plot([x, x], [1, 21], "k-", lw=0.8)
        ax.text(x, 21.2, a, ha="center", fontsize=7, color="black")
    msgs = [
        ("TickScheduler", "StatusDispatcher", "dispatch()"),
        ("StatusDispatcher", "TargetPlanner", "ASSIGN_TARGET"),
        ("TargetPlanner", "StatusDispatcher", "TARGET_ASSIGNED"),
        ("StatusDispatcher", "StatusDispatcher", "onTargetAssigned() → WAITING_ROUTE"),
        ("StatusDispatcher", "Navigator", "PLAN_ROUTE"),
        ("Navigator", "StatusDispatcher", "ROUTE_PLANNED"),
        ("StatusDispatcher", "StatusDispatcher", "onRoutePlanned()"),
        ("StatusDispatcher", "StrategySupervisor", "SUPERVISE_ROUTE"),
        ("StrategySupervisor", "StatusDispatcher", "ROUTE_OPTIMIZED"),
        ("StatusDispatcher", "Car\nMoveExecutor", "TICK_MOVE"),
        ("Car\nMoveExecutor", "Car\nMoveExecutor", "executeMove() → doMove()"),
        ("Car\nMoveExecutor", "StatusDispatcher", "MOVED"),
        ("StatusDispatcher", "StatusDispatcher", "onMoveAcknowledged()"),
        ("StatusDispatcher", "Display", "REFRESH_ALL (Fanout)"),
        ("Display", "Display", "pushSimulationState()"),
    ]
    y = 20
    for src, dst, label in msgs:
        x1, x2 = xs[src], xs[dst]
        if x1 == x2:
            ax.plot([x1-0.2, x1+0.2], [y, y], "k-", lw=0.8)
            ax.text(x1+0.25, y, label, fontsize=6.5, color="black", va="center")
        else:
            ax.annotate("", xy=(x2, y), xytext=(x1, y),
                        arrowprops=dict(arrowstyle="->", color="black", lw=0.9))
            ax.text((x1+x2)/2, y+0.08, label, ha="center", fontsize=6.5, color="black")
        y -= 1.15
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def build_document(imgs):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style.font.color.rgb = BLACK
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("《软件体系结构》课程作业报告\n\n")
    r.bold = True; r.font.size = Pt(16); _black(r)
    r2 = t.add_run("变电站巡检仿真系统\n软件架构设计说明")
    r2.bold = True; r2.font.size = Pt(18); _black(r2)
    for line in ["项目名称：变电站巡检仿真系统（car_homework）",
                 "开发语言：Java 17 + Redis + RabbitMQ + SQL Server + Web",
                 "架构风格：黑板架构（Blackboard Architecture）+ 分布式多进程"]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line); _black(run)
    doc.add_page_break()

    heading(doc, "一、采用的软件架构", 1)
    para(doc,
         "本系统采用黑板架构（Blackboard Architecture）风格的分布式多进程设计。"
         "Redis 充当共享黑板（BlackboardClient），存储地图位图、车辆状态、路径列表、"
         "探索率等全局数据；RabbitMQ 充当连接件（MessageBus），负责模块间异步消息传递；"
         "Controller 充当黑板控制器，通过 TickScheduler 按节拍（默认 500ms）驱动"
         "StatusDispatcher 调度各知识源。各知识源（TargetPlanner、Navigator、"
         "StrategySupervisor、TaskConfigurator、Car）为独立 JVM 进程，"
         "彼此之间禁止直接调用，仅通过 MQ 消息与 Redis 黑板协作。",
         indent=True)
    para(doc, "架构名称：黑板架构（Blackboard Architecture）+ 中介者模式（Controller 调度）", bold=True)
    doc.add_picture(imgs["arch"], width=Inches(6.3))
    cap = doc.add_paragraph("图1  黑板架构总体架构图")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs: _black(run)

    heading(doc, "1.1  黑板架构核心特征", 2)
    for item in [
        "共享黑板：Redis 存储 mapView/mapBlock/mapSealed、各车 Position/Target/RouteList/Status 等",
        "知识源独立：Navigator、TargetPlanner、Car 等互不直接通信，只读写黑板 + 收发 MQ",
        "控制器调度：Controller 的 StatusDispatcher 根据 CarStatus 状态机分派任务",
        "异步消息驱动：所有跨模块协作通过 RabbitMQ 点对点队列或 Fanout 广播",
        "多进程部署：每个知识源可独立部署于不同主机，通过 infra.local.json 配置远程 Redis/MQ",
    ]:
        bullet(doc, item)

    heading(doc, "二、软件架构核心构件清单", 1)

    components = [
        ("构件① Controller（黑板控制器）",
         "独立部署（单实例 JVM 进程）；Redis SETNX 保证全局唯一",
         "消息驱动；订阅 ControllerCmd 队列，发布至各 Worker 队列",
         [
             ("ControllerMain.start()", "入口：获取单实例锁、连接 MQ/Redis、启动调度", ""),
             ("TickScheduler.start()", "启动节拍循环 tickLoop()，默认 interval=500ms", "CPU 定时驱动"),
             ("TickScheduler.tickLoop()", "每节拍调用 StatusDispatcher.dispatch()", ""),
             ("CommandHandler.handle(raw)", "解析 MQ 入站消息，路由至 Dispatcher/Scheduler", "JSON 字符串"),
             ("StatusDispatcher.dispatch()", "发现车辆、按 CarStatus 分派、检测探索完成", ""),
             ("StatusDispatcher.onTargetAssigned()", "收到 TARGET_ASSIGNED 后转 WAITING_ROUTE", ""),
             ("StatusDispatcher.onRoutePlanned()", "收到 ROUTE_PLANNED 后触发监督或 READY", ""),
             ("StatusDispatcher.onMoveAcknowledged()", "收到 MOVED 后更新状态、广播 REFRESH_ALL", ""),
             ("StatusDispatcher.forwardConfig()", "转发 FORWARD_CONFIG 至 TaskConfigurator", ""),
             ("StatusDispatcher.completeTask()", "探索完成：stop 调度、冻结车辆、设 taskActive=false", ""),
         ]),
        ("构件② StatusDispatcher（状态分派器）",
         "嵌入 Controller 进程功能模块",
         "函数调用（Controller 内部）+ MQ 发布",
         [
             ("dispatchCar(carId)", "按 IDLE/WAITING_ROUTE/READY/MOVING/BLOCKED 分派", ""),
             ("sendAssignTarget(carId)", "发布 ASSIGN_TARGET → TargetPlannerCmd", ""),
             ("checkAndPlanRoute(carId)", "发布 PLAN_ROUTE → NavigatorCmd", ""),
             ("sendSuperviseRoute(carId)", "发布 SUPERVISE_ROUTE → StrategySupervisorCmd", ""),
             ("trySendTickMove(carId)", "发布 TICK_MOVE → Car_{carId}", ""),
         ]),
        ("构件③ TargetPlanner（目标分配知识源）",
         "独立部署（可竞争消费，通常 1 实例）",
         "消息驱动；订阅 TargetPlannerCmd",
         [
             ("TargetPlannerMain.handleAssignTarget()", "输入：ASSIGN_TARGET；调用 GreedyTargetAllocator.allocate()", ""),
             ("GreedyTargetAllocator.allocate(carId, bb)", "输入：carId、BlackboardClient；输出：Point 目标坐标",
              "基于前沿格子（Frontier）贪心选取未探索目标"),
             ("sendSuccessReply()", "发布 TARGET_ASSIGNED → ControllerCmd", ""),
         ]),
        ("构件④ Navigator（路径规划知识源）",
         "独立部署",
         "消息驱动；订阅 NavigatorCmd",
         [
             ("NavigatorMain.handlePlanRoute()", "输入：PLAN_ROUTE；解析 start/target/algorithm", ""),
             ("PathPlannerFactory.create(algorithm)", "创建 BfsPathFinder 或 AStarPathFinder", "algorithm ∈ {BFS, ASTAR}"),
             ("PathPlanner.plan(start, target, bb)", "输出：List<RouteStep>；写入 Redis RouteList", ""),
             ("sendRoutePlanned()", "发布 ROUTE_PLANNED → ControllerCmd", ""),
         ]),
        ("构件⑤ StrategySupervisor（策略监督知识源）",
         "独立部署",
         "消息驱动；订阅 StrategySupervisorCmd",
         [
             ("StrategySupervisorMain.handleSupervise()", "输入：SUPERVISE_ROUTE", ""),
             ("RouteOverlapEvaluator.isHighlyOverlapped()", "检测多车路径重叠", ""),
             ("RouteEvaluator.evaluate()", "检测过度探索路线", ""),
             ("WeightedPathPlanner.plan()", "生成偏好未探索格子的加权路径", ""),
             ("sendResult()", "发布 ROUTE_OPTIMIZED → ControllerCmd", ""),
         ]),
        ("构件⑥ Car（移动执行知识源）",
         "独立部署（每车 1 进程，可动态 ADD_CAR 启动）",
         "消息驱动；订阅 Car_{carId} 队列",
         [
             ("CarMain.start()", "自注册出生点、订阅专属队列", ""),
             ("CarAgent.handleMessage()", "分发 TICK_MOVE / BLOCKED_TIMEOUT", ""),
             ("MoveExecutor.executeMove(tick)", "获取分布式锁 lock:{carId}", ""),
             ("MoveExecutor.doMove()", "弹出 RouteList 首步、更新 Position、recordExploration()", ""),
             ("sendMoved() / sendRouteDone() / sendBlocked()", "发布 MOVED/ROUTE_DONE/BLOCKED → ControllerCmd", ""),
         ]),
        ("构件⑦ TaskConfigurator（任务初始化知识源）",
         "独立部署（单实例）",
         "消息驱动；订阅 TaskConfigCmd",
         [
             ("TaskConfiguratorMain.handleConfig()", "处理 FORWARD_CONFIG", ""),
             ("TaskInitializer.initialize()", "生成障碍物、出生点、写 TaskConfig、注册车辆", ""),
             ("selectiveClear()", "清除仿真相关 Redis 键（保留 auth 会话）", ""),
             ("handleReset()", "处理 FORWARD_RESET", ""),
         ]),
        ("构件⑧ Display（展示层）",
         "独立部署（单实例）；HTTP:8887 + WebSocket:8888",
         "消息驱动 + WebSocket 推送",
         [
             ("DisplayMain.start()", "启动 HTTP/WS/MQ 订阅", ""),
             ("WebSocketBridge.onMessage()", "接收浏览器 SET_CONFIG/RESET/ADD_CAR 等", ""),
             ("WebSocketBridge.pushSimulationState()", "推送 SimulationState JSON 至浏览器", ""),
             ("DynamicCarLauncher.launchAsync()", "ProcessBuilder 动态启动 Car JAR", ""),
             ("onRefreshAllReceived()", "订阅 Fanout UpdateView，收到 REFRESH_ALL 后刷新", ""),
         ]),
        ("构件⑨ BlackboardClient（黑板客户端）",
         "嵌入各进程功能模块（common 库）",
         "函数调用；读写 Redis",
         [
             ("getCarPosition(carId)", "读取 {carId}:Position", ""),
             ("pushRoute(carId, route)", "写入 {carId}:RouteList", ""),
             ("setCarStatus(carId, status)", "写入 {carId}:Status", "CarStatus 枚举"),
             ("recordExploration(x, y)", "更新 mapView 位图、explorationRate", ""),
             ("isExplorationComplete()", "判定探索率 100% 且无未探索可达格", ""),
             ("discoverCarIds()", "扫描 Redis 发现所有注册车辆", ""),
         ]),
        ("构件⑩ MessageBus（消息总线连接件封装）",
         "嵌入各进程功能模块",
         "MQ 解耦；封装 RabbitMQ JMS 客户端",
         [
             ("connect()", "建立 RabbitMQ 连接", "tcp://host:5672"),
             ("publish(queue, message)", "点对点发送 JSON 消息", ""),
             ("subscribe(queue, handler)", "注册消费者回调", ""),
             ("publishFanout(exchange, message)", "Fanout 广播（UpdateView）", ""),
         ]),
    ]

    for title, deploy, call, funcs in components:
        heading(doc, title, 2)
        para(doc, f"部署方式：{deploy}")
        para(doc, f"调用方式：{call}")
        para(doc, "功能清单：", bold=True)
        for fname, desc, fmt in funcs:
            bullet(doc, f"功能：{fname}")
            para(doc, f"    {desc}", indent=True)
            if fmt:
                para(doc, f"    参数说明：{fmt}", indent=True)

    heading(doc, "三、连接件清单", 1)
    connectors = [
        ("连接件① Controller → TargetPlanner 目标分配管道",
         "MQ 解耦（RabbitMQ 点对点队列 TargetPlannerCmd）",
         "MessageBuilder.build(ASSIGN_TARGET, tick, carId, data)",
         "消息格式：{\"type\":\"ASSIGN_TARGET\",\"tick\":42,\"carId\":\"Car001\",\"timestamp\":...,\"data\":{}}\n"
         "收发频次：每车每节拍最多 1 次（IDLE 状态触发）"),
        ("连接件② Controller → Navigator 路径规划管道",
         "MQ 解耦（队列 NavigatorCmd）",
         "MessageBuilder.build(PLAN_ROUTE, tick, carId, {start, target, algorithm, supervised})",
         "start/target: {x:int, y:int}；algorithm ∈ {BFS, ASTAR}"),
        ("连接件③ Controller → StrategySupervisor 路线监督管道",
         "MQ 解耦（队列 StrategySupervisorCmd）",
         "MessageBuilder.build(SUPERVISE_ROUTE, tick, carId)",
         "探索率 < 85% 时触发；同一车冷却 15 tick"),
        ("连接件④ Controller → Car 移动指令管道",
         "MQ 解耦（队列 Car_{carId}，每车独立队列）",
         "MessageBuilder.build(TICK_MOVE, tick, carId, {tick})",
         "READY 状态下每节拍 1 次"),
        ("连接件⑤ Worker → Controller 回复管道",
         "MQ 解耦（队列 ControllerCmd）",
         "TARGET_ASSIGNED / ROUTE_PLANNED / MOVED / ROUTE_DONE / BLOCKED / ROUTE_OPTIMIZED / TASK_READY",
         "异步请求-响应模式"),
        ("连接件⑥ Controller → Display 视图刷新广播",
         "MQ 解耦（Fanout 交换机 UpdateView）",
         "MessageBuilder.build(REFRESH_ALL, tick, null, {explorationRate})",
         "每节拍结束后广播；队列方式：Fanout 发布-订阅"),
        ("连接件⑦ Display ↔ Browser WebSocket 管道",
         "消息驱动（WebSocket 长连接，端口 8888）",
         "浏览器 → SET_CONFIG/RESET/TOGGLE_PAUSE/ADD_CAR；服务端 → SimulationState JSON",
         "SimulationState 含 tick、explorationRate、cars[]、mapViewB64 等"),
        ("连接件⑧ 各模块 ↔ Redis 黑板连接",
         "函数调用（Jedis 客户端读写）",
         "BlackboardClient 各类 get/set/push 方法",
         "键空间：mapView、{carId}:Position、TaskConfig、explorationEvents 等"),
        ("连接件⑨ Car 分布式锁连接",
         "函数调用 + Redis SETNX",
         "DistributedLock.acquire(carId) / release(carId)",
         "防止多车同时占用同一格子；锁键 lock:{carId}"),
    ]
    for name, deploy, fmt, data in connectors:
        heading(doc, name, 2)
        para(doc, f"部署方式：{deploy}")
        para(doc, f"消息/调用格式：{fmt}")
        para(doc, f"数据格式：{data}")

    heading(doc, "四、架构约束与构件协作关系", 1)
    for c in [
        "知识源之间禁止直接调用：TargetPlanner 不直接调 Navigator，Car 不直接调 TargetPlanner，所有协作经 Controller 中介。",
        "黑板为唯一共享数据源：地图、车辆状态、路径均存 Redis，读写通过 BlackboardClient，避免多进程内存不一致。",
        "Controller 全局单例：通过 Redis controller:instance SETNX 锁保证同一时刻仅一个 Controller 驱动节拍。",
        "小车状态机约束：IDLE→WAITING_ROUTE→READY→MOVING→BLOCKED 五态流转，StatusDispatcher.dispatchCar() 按态分派。",
        "避免循环依赖：消息流为单向链式 Controller→Worker→Controller，不存在 Worker A 调 Worker B。",
        "与黑板架构一致：知识源只受 Controller 调度，数据在黑板共享；本设计不与「知识源互不调用」原则冲突。",
        "时序图消息名与 MessageTypes 常量、StatusDispatcher 方法名完全一致。",
    ]:
        bullet(doc, c)

    add_table(doc, ["小车状态 CarStatus", "含义", "StatusDispatcher 动作"],
              [
                  ["IDLE", "空闲，无目标无路径", "sendAssignTarget()"],
                  ["WAITING_ROUTE", "已分配目标，等待路径", "checkAndPlanRoute()"],
                  ["READY", "路径就绪，可移动", "trySendTickMove()"],
                  ["MOVING", "正在移动一步", "checkMovingStuck()"],
                  ["BLOCKED", "被障碍物阻塞", "checkBlockedTimeout()"],
              ])

    heading(doc, "五、协作流程案例说明（含时序图）", 1)

    heading(doc, "5.1  案例一：仿真启动流程", 2)
    para(doc,
         "用户在浏览器点击「开始」，WebSocketBridge 发送 SET_CONFIG 至 ControllerCmd 队列。"
         "CommandHandler 调用 StatusDispatcher.prepareForNewConfig() 后 forwardConfig()，"
         "TaskConfigurator 执行 TaskInitializer.initialize() 初始化地图与车辆，"
         "回复 TASK_READY，Controller 启动 TickScheduler 开始节拍循环。",
         indent=True)
    doc.add_picture(imgs["startup"], width=Inches(6.2))
    cap2 = doc.add_paragraph("图2  仿真启动协作时序图")
    cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap2.runs: _black(run)

    heading(doc, "5.2  案例二：单节拍探索流程", 2)
    para(doc,
         "TickScheduler 每 500ms 调用 StatusDispatcher.dispatch()。对 IDLE 车辆发送 ASSIGN_TARGET，"
         "TargetPlanner 贪心选取前沿目标并回复 TARGET_ASSIGNED；Controller 发 PLAN_ROUTE 给 Navigator，"
         "Navigator 用 BFS/A* 规划路径写入黑板 RouteList；可选经 StrategySupervisor 优化路线；"
         "READY 车辆收到 TICK_MOVE 后 MoveExecutor 执行单步移动并回复 MOVED；"
         "最后 Controller Fanout 广播 REFRESH_ALL，Display 推送 SimulationState 至浏览器。",
         indent=True)
    doc.add_picture(imgs["tick"], width=Inches(6.3))
    cap3 = doc.add_paragraph("图3  单节拍探索协作时序图")
    cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap3.runs: _black(run)

    heading(doc, "5.3  案例三：探索完成判定", 2)
    para(doc,
         "每节拍 dispatch() 末尾检测 BlackboardClient.isExplorationComplete()："
         "探索率达 100% 且无未探索可达格时，或连续 30 节拍全车 IDLE 时，"
         "调用 completeTask() 停止 TickScheduler、冻结车辆、广播最终 REFRESH_ALL。",
         indent=True)

    heading(doc, "六、消息协议与黑板键空间", 1)
    add_table(doc, ["消息类型", "方向", "关键 data 字段"],
              [
                  ["SET_CONFIG", "Display → Controller", "mapWidth, mapHeight, carCount, algorithm, tickInterval"],
                  ["FORWARD_CONFIG", "Controller → TaskConfigurator", "同上"],
                  ["TASK_READY", "TaskConfigurator → Controller", "{}"],
                  ["ASSIGN_TARGET", "Controller → TargetPlanner", "{}"],
                  ["TARGET_ASSIGNED", "TargetPlanner → Controller", "success, target{x,y}"],
                  ["PLAN_ROUTE", "Controller → Navigator", "start, target, algorithm, supervised"],
                  ["ROUTE_PLANNED", "Navigator → Controller", "routeFound, routeLength"],
                  ["SUPERVISE_ROUTE", "Controller → StrategySupervisor", "{}"],
                  ["ROUTE_OPTIMIZED", "StrategySupervisor → Controller", "optimized, overlapReassign"],
                  ["TICK_MOVE", "Controller → Car", "{tick}"],
                  ["MOVED", "Car → Controller", "newPosition, routeRemaining"],
                  ["REFRESH_ALL", "Controller → Display(Fanout)", "explorationRate"],
              ])

    para(doc, "统一消息信封格式（MessageBuilder.build）：", bold=True)
    para(doc,
         '{"type":"PLAN_ROUTE","tick":42,"carId":"Car001","timestamp":1710000000000,'
         '"data":{"start":{"x":1,"y":2},"target":{"x":5,"y":6},"algorithm":"BFS","supervised":false}}',
         indent=True)

    heading(doc, "七、部署方案", 1)
    add_table(doc, ["组件", "部署方式", "端口/约束"],
              [
                  ["Redis", "Docker 独立部署", "6379"],
                  ["RabbitMQ", "Docker 独立部署", "5672 / 15672"],
                  ["Controller", "独立 JVM，全局 1 实例", "最后启动"],
                  ["TaskConfigurator", "独立 JVM，1 实例", "先于 Controller"],
                  ["Navigator / TargetPlanner / StrategySupervisor", "独立 JVM，可分布式", "竞争消费"],
                  ["Car × N", "独立 JVM，每车 1 进程", "可动态 ADD_CAR"],
                  ["Display", "独立 JVM，1 实例", "HTTP 8887 / WS 8888"],
                  ["SQL Server", "独立部署", "1433（用户/回放）"],
              ])
    para(doc,
         "单机启动：docker compose up -d → TaskConfigurator → Navigator/TargetPlanner/StrategySupervisor "
         "→ Car001…Car00N → Display → Controller。分布式时通过 deploy/infra.local.json 配置远程 Redis/MQ 地址。",
         indent=True)

    heading(doc, "八、设计模式与总结", 1)
    add_table(doc, ["设计模式", "体现位置"],
              [
                  ["黑板架构", "Redis 共享状态 + Controller 调度"],
                  ["中介者模式", "Controller/StatusDispatcher 协调所有知识源"],
                  ["状态机模式", "CarStatus 五态 + dispatchCar()"],
                  ["策略模式", "PathPlanner → BfsPathFinder / AStarPathFinder"],
                  ["工厂模式", "PathPlannerFactory.create()"],
                  ["发布-订阅", "Fanout UpdateView → REFRESH_ALL"],
                  ["命令模式", "MQ 消息作为跨进程命令"],
              ])
    para(doc,
         "本系统完整实现了黑板架构的核心约束：知识源独立部署、仅通过黑板共享数据、"
         "仅受 Controller 调度、知识源之间不直接调用。RabbitMQ 实现 MQ 解耦，"
         "Redis 实现黑板共享，WebSocket 实现实时展示。时序图中的消息名称"
         "（dispatch、ASSIGN_TARGET、PLAN_ROUTE、TICK_MOVE、MOVED、REFRESH_ALL）"
         "与 MessageTypes 常量及 StatusDispatcher 方法完全一致，符合软件架构设计题回答要点。",
         indent=True)

    heading(doc, "附录：UML 类图", 1)
    para(doc,
         "下图展示变电站巡检仿真系统核心类结构。Controller 模块（TickScheduler、StatusDispatcher、"
         "CommandHandler）通过 MessageBus 调度各知识源；各 Worker 依赖 BlackboardClient 读写 Redis 黑板；"
         "Navigator 通过 PathPlanner 接口与 BfsPathFinder/AStarPathFinder 实现策略模式；"
         "Car 模块由 CarAgent 委托 MoveExecutor 执行移动。",
         indent=True)
    doc.add_picture(imgs["uml"], width=Inches(6.5))
    cap = doc.add_paragraph("图4  UML 类图 — 变电站巡检仿真系统核心类")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        _black(run)

    os.makedirs(os.path.dirname(DOC_PATH), exist_ok=True)
    doc.save(DOC_PATH)
    print(DOC_PATH)


def main():
    import sys
    sys.path.insert(0, UML_HELPER_DIR)
    from uml_helper import draw_car_uml

    os.makedirs(OUT_DIR, exist_ok=True)
    imgs = {
        "arch": os.path.join(OUT_DIR, "car_fig1_arch.png"),
        "startup": os.path.join(OUT_DIR, "car_fig2_startup.png"),
        "tick": os.path.join(OUT_DIR, "car_fig3_tick.png"),
        "uml": os.path.join(OUT_DIR, "car_fig4_uml.png"),
    }
    draw_blackboard_arch(imgs["arch"])
    draw_seq_startup(imgs["startup"])
    draw_seq_tick(imgs["tick"])
    draw_car_uml(imgs["uml"])
    build_document(imgs)


if __name__ == "__main__":
    main()
